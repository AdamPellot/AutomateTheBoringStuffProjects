#! python3
# customInvite.py - Generates a Word document with custom invitations.
# Adam Pellot

import docx
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

print('Enter the absolute path of the blank word doc with your styles:')
filePath = input()
doc = docx.Document(filePath)

print('Enter the absolute path of the file with your guest list')
guestsPath = input()
guestsFile = open(guestsPath)

for name in guestsFile.readlines():
        name = name.rstrip('\n')
        p1 = doc.add_paragraph('It would be a pleasure to have the company of')
        font = p1.runs[0].font
        font.name = 'Brush Script Std'
        font.size = Pt(36)
        p1.alignment = 1
        p2 = doc.add_paragraph(name)
        p2.runs[0].bold = True
        font = p2.runs[0].font
        font.name = 'Times New Roman'
        font.size = Pt(36)
        p2.alignment = 1
        p3 = doc.add_paragraph('In Wayne Manor on the evening of')
        font = p3.runs[0].font
        font.name = 'Brush Script Std'
        font.size = Pt(36)
        p3.alignment = 1
        p4 = doc.add_paragraph('April 7th')
        font = p4.runs[0].font
        font.name = 'Times New Roman'
        font.size = Pt(36)
        p4.alignment = 1
        p5 = doc.add_paragraph("At 7 o'clock")
        font = p5.runs[0].font
        font.name = 'Brush Script Std'
        font.size = Pt(36)
        p5.alignment = 1

        doc.add_page_break()

print('Enter the absolute path of where you would like the invitations saved:')
finalPath = input()

doc.save(finalPath)
